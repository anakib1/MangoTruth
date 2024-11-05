from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def extract_docx_data(file):
    text_data = []
    table_data = []
    try:
        doc = Document(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"<'{file}' was not found>")
    except PackageNotFoundError:
        raise ValueError(f"<'{file}' is not a .docx file>")

    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_data.append(paragraph.text)

    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_content.append(row_data)
        table_data.append(table_content)

    return text_data, table_data